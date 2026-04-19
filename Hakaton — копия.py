import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import os
import re
import csv
from pathlib import Path
from datetime import datetime
import sys
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
from functools import lru_cache
import hashlib
import json
from collections import defaultdict

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


class FastPersonalDataDetector:
    def __init__(self):
        # Компилируем регулярки один раз для скорости
        self.patterns = {
            'snils': re.compile(r'\b\d{3}[\-\s]?\d{3}[\-\s]?\d{3}[\-\s]?\d{2}\b'),
            'passport': re.compile(r'\b\d{4}[\-\s]?\d{6}\b'),
            'inn': re.compile(r'\b\d{12}\b'),
            'phone': re.compile(r'\b(?:8|\+7)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}\b'),
            'email': re.compile(r'\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b'),
        }

    def has_personal_data(self, text):
        if not text or len(text) < 10:
            return False

        # Быстрая проверка - ищем хотя бы одну цифру или @
        if not any(c.isdigit() or c == '@' for c in text[:500]):
            return False

        found = []
        for name, pattern in self.patterns.items():
            # Ищем только первые 3 совпадения для скорости
            matches = pattern.findall(text)
            if matches:
                found.append(f"{name}: {matches[:2]}")

        return len(found) > 0


class FastImageProcessor:
    @staticmethod
    def quick_preprocess(image):
        """Быстрая предобработка для OCR"""
        try:
            # Конвертируем в оттенки серого если нужно
            if image.mode != 'L':
                image = image.convert('L')

            # Оптимальное увеличение - 2x вместо 3x
            if image.width < 1000:
                new_size = (image.width * 2, image.height * 2)
                image = image.resize(new_size, Image.Resampling.LANCZOS)

            # Только базовая обработка для скорости
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(1.5)

            return image
        except:
            return image


class FileProcessorFast:
    def __init__(self, max_workers=4):
        self.detector = FastPersonalDataDetector()
        self.image_processor = FastImageProcessor()
        self.max_workers = max_workers
        self.cache_file = "ocr_cache.json"
        self.cache = self.load_cache()

    def load_cache(self):
        """Загрузка кэша OCR результатов"""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def save_cache(self):
        """Сохранение кэша"""
        try:
            # Ограничиваем размер кэша 1000 записей
            if len(self.cache) > 1000:
                # Оставляем только последние 1000
                self.cache = dict(list(self.cache.items())[-1000:])
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f)
        except:
            pass

    def get_file_hash(self, filepath):
        """Быстрый хэш файла для кэширования"""
        try:
            with open(filepath, 'rb') as f:
                # Читаем только первые 64KB для скорости
                return hashlib.md5(f.read(65536)).hexdigest()
        except:
            return None

    def find_files_fast(self, directory):
        """Быстрый поиск файлов с использованием генератора"""
        directory_path = Path(directory)

        # Расширения в нижнем регистре для быстрого сравнения
        extensions = {'.pdf', '.docx', '.xlsx', '.tif', '.tiff',
                      '.jpg', '.jpeg', '.png'}

        # Используем rglob для рекурсивного поиска
        for file_path in directory_path.rglob('*'):
            if file_path.suffix.lower() in extensions:
                yield file_path

    def extract_text_from_image_fast(self, filepath, lang='rus+eng'):
        """Быстрое извлечение текста из изображений"""
        try:
            # Проверяем кэш
            file_hash = self.get_file_hash(filepath)
            if file_hash and file_hash in self.cache:
                return self.cache[file_hash]

            # Оптимизированная обработка
            with Image.open(filepath) as img:
                # Для маленьких изображений не увеличиваем
                if img.width < 500 or img.height < 500:
                    processed = self.image_processor.quick_preprocess(img)
                    config = '--oem 3 --psm 6'  # Только одна конфигурация
                else:
                    # Для больших изображений только базовая обработка
                    processed = img.convert('L') if img.mode != 'L' else img
                    config = '--oem 3 --psm 6'

                text = pytesseract.image_to_string(processed, lang=lang, config=config)
                text = re.sub(r'\s+', ' ', text).strip()

                # Сохраняем в кэш
                if file_hash and len(text) > 50:
                    self.cache[file_hash] = text

                return text

        except Exception as e:
            return ""

    def extract_text_from_pdf_fast(self, filepath):
        """Быстрое извлечение текста из PDF"""
        try:
            import fitz

            text = ""
            doc = fitz.open(filepath)

            # Ограничиваем количество страниц для больших PDF
            max_pages = min(len(doc), 50)  # Максимум 50 страниц
            pages_processed = 0

            for page_num in range(max_pages):
                page = doc[page_num]
                page_text = page.get_text()

                if page_text.strip():
                    text += page_text[:2000] + " "  # Ограничиваем текст на странице
                elif pages_processed < 50:  # OCR только для первых 10 страниц
                    # Уменьшаем разрешение для OCR
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    ocr_text = self.extract_text_from_image_fast(img)
                    if ocr_text:
                        text += ocr_text[:1000] + " "
                    pages_processed += 1

                # Ранний выход если уже нашли данные
                if len(text) > 5000:
                    break

            doc.close()
            return text[:10000]  # Ограничиваем общий размер текста

        except Exception as e:
            return ""

    def extract_text_from_docx_fast(self, filepath):
        """Быстрое извлечение текста из DOCX"""
        try:
            from docx import Document

            text = []
            doc = Document(filepath)

            # Быстрый сбор текста
            for paragraph in doc.paragraphs[:100]:  # Ограничиваем количество параграфов
                if paragraph.text:
                    text.append(paragraph.text[:1000])  # Ограничиваем длину параграфа

            # Проверяем только первые 2 таблицы
            for table in doc.tables[:10]:
                for row in table.rows[:100]:  # Ограничиваем строки в таблице
                    for cell in row.cells[:50]:  # Ограничиваем ячейки
                        if cell.text:
                            text.append(cell.text[:200])

            return ' '.join(text)[:5000]

        except Exception as e:
            return ""

    def extract_text_from_excel_fast(self, filepath):
        """Быстрое извлечение текста из Excel"""
        try:
            import pandas as pd

            text = []
            # Читаем только первые 2 листа
            excel_file = pd.ExcelFile(filepath)
            for sheet_name in excel_file.sheet_names[:6]:
                # Читаем только первые 100 строк
                df = pd.read_excel(filepath, sheet_name=sheet_name, nrows=100)
                # Конвертируем только строковые столбцы
                for col in df.select_dtypes(include=['object']).columns:
                    text.extend(df[col].astype(str).head(50).tolist())

            return ' '.join(text)[:5000]

        except Exception as e:
            return ""

    def process_file(self, filepath):
        """Быстрая обработка одного файла"""
        try:
            file_path_obj = Path(filepath)
            file_ext = file_path_obj.suffix.lower()
            text = ""

            # Выбираем метод в зависимости от расширения
            if file_ext == '.pdf':
                text = self.extract_text_from_pdf_fast(filepath)
            elif file_ext == '.docx':
                text = self.extract_text_from_docx_fast(filepath)
            elif file_ext == '.xlsx':
                text = self.extract_text_from_excel_fast(filepath)
            elif file_ext in ['.tif', '.tiff', '.jpg', '.jpeg', '.png']:
                text = self.extract_text_from_image_fast(filepath)

            if text and len(text) > 20:
                if self.detector.has_personal_data(text):
                    size = os.path.getsize(filepath)
                    mod_time = datetime.fromtimestamp(os.path.getmtime(filepath))
                    return {
                        'size': size,
                        'time': mod_time.strftime("%Y-%m-%d %H:%M:%S"),
                        'name': str(file_path_obj.name),
                        'full_path': str(filepath)
                    }

        except Exception as e:
            pass
        return None

    def process_file_wrapper(self, args):
        """Обертка для многопоточности"""
        filepath, = args
        return self.process_file(filepath)

    def process_directory(self, input_dir, output_csv, max_workers=None):
        """Быстрая обработка директории с параллелизацией"""
        print(f"Scanning directory: {input_dir}")

        # Быстрый сбор файлов
        files = list(self.find_files_fast(input_dir))

        if not files:
            print("No supported files found")
            return

        total = len(files)
        print(f"Found {total} files")

        # Используем максимальное количество потоков
        workers = max_workers or min(self.max_workers, os.cpu_count() or 1)
        print(f"Using {workers} workers for parallel processing\n")

        results = []
        processed = 0

        # Параллельная обработка с ThreadPoolExecutor
        with ThreadPoolExecutor(max_workers=workers) as executor:
            # Отправляем все задачи
            futures = {executor.submit(self.process_file, str(f)): f for f in files}

            for future in futures:
                processed += 1
                result = future.result()
                if result:
                    results.append(result)
                    print(f"\n✓ [{processed}/{total}] FOUND PERSONAL DATA in {result['name']}")
                else:
                    # Показываем прогресс
                    if processed % 1 == 0:
                        print(f"Processed {processed}/{total} files...")

        # Сохраняем кэш
        self.save_cache()

        # Сохраняем результаты
        with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
            writer = csv.DictWriter(csvfile, fieldnames=['size', 'time', 'name', 'full_path'])
            writer.writeheader()
            writer.writerows(results)

        print(f"\n{'=' * 60}")
        print(f"Results saved to: {output_csv}")
        print(f"Files with personal data: {len(results)}/{total}")
        print(f"Time saved: {len(results)} files with personal data found")
        print(f"{'=' * 60}")


# Более быстрая версия с минимальной обработкой
class UltraFastFileProcessor(FileProcessorFast):
    def extract_text_from_image_ultrafast(self, filepath, lang='rus+eng'):
        """Максимально быстрое извлечение текста"""
        try:
            # Проверяем кэш
            file_hash = self.get_file_hash(filepath)
            if file_hash and file_hash in self.cache:
                return self.cache[file_hash]

            # Минимальная обработка
            with Image.open(filepath) as img:
                # Только конвертация в серый
                if img.mode != 'L':
                    img = img.convert('L')

                # Простое распознавание
                text = pytesseract.image_to_string(img, lang=lang, config='--psm 6')
                text = text.replace('\n', ' ').strip()

                # Кэшируем только если есть цифры
                if file_hash and any(c.isdigit() for c in text):
                    self.cache[file_hash] = text

                return text[:2000]  # Ограничиваем размер
        except:
            return ""

    def process_file(self, filepath, lang='rus+eng'):
        try:
            # Преобразуем строку в Path объект
            file_path_obj = Path(filepath)
            file_ext = file_path_obj.suffix.lower()
            text = ""

            print(f"  Processing {file_path_obj.name}...")  # Используем Path объект

            if file_ext == '.pdf':
                text = self.extract_text_from_pdf(filepath)
            elif file_ext == '.docx':
                text = self.extract_text_from_docx(filepath)
            elif file_ext in ['.xls', '.xlsx']:
                text = self.extract_text_from_excel(filepath)
            elif file_ext in ['.tif', '.tiff', '.jpg', '.jpeg', '.png']:
                text = self.extract_text_from_image(filepath, lang)

            if text and len(text.strip()) > 20:
                print(f"  Text extracted: {len(text)} chars")
                print(f"  Preview: {text[:200]}...")

                if self.detector.has_personal_data(text):
                    size, time_str = self.get_file_info(filepath)
                    return {
                        'size': size,
                        'time': time_str,
                        'name': str(file_path_obj.name),
                        'full_path': str(filepath),
                        'text_preview': text[:500]
                    }
                else:
                    print(f"  No personal data patterns found")
            else:
                print(f"  Failed to extract text or text too short ({len(text) if text else 0} chars)")

        except Exception as e:
            print(f"  Error: {e}")
            import traceback
            traceback.print_exc()
        return None


if __name__ == "__main__":
    input_dir = sys.argv[1] if len(sys.argv) > 1 else "./"
    output_csv = sys.argv[2] if len(sys.argv) > 2 else "result.csv"

    processor = FileProcessorFast(max_workers=4)
    processor.process_directory(input_dir, output_csv)