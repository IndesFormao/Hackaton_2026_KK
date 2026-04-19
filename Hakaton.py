import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import os
import re
import csv
from pathlib import Path
from datetime import datetime
import sys
from concurrent.futures import ThreadPoolExecutor, as_completed
import hashlib
import json
from collections import defaultdict
from difflib import SequenceMatcher
import numpy as np

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


class AdvancedPersonalDataDetector:
    def __init__(self):
        # Расширенные паттерны с улучшенной точностью
        self.patterns = {
            'snils': [
                re.compile(r'\b(\d{3}[\-\s]?\d{3}[\-\s]?\d{3}[\-\s]?\d{2})\b'),
                re.compile(r'(?:СНИЛС|страховой\s+номер|snils)[:\s=]*(\d{3}[\-\s]?\d{3}[\-\s]?\d{3}[\-\s]?\d{2})',
                           re.IGNORECASE)
            ],
            'passport': [
                re.compile(r'\b(\d{2}\s*\d{2}[\s\-]?\d{6})\b'),
                re.compile(r'\b(\d{4}[\s\-]?\d{6})\b'),
                re.compile(
                    r'(?:паспорт|passport|passport\s+no|серия\s+и\s+номер)[:\s=]*(\d{2}[\s\-]?\d{2}[\s\-]?\d{6})',
                    re.IGNORECASE)
            ],
            'inn': [
                re.compile(r'\b(\d{12})\b'),
                re.compile(r'\b(\d{10})\b'),  # ИНН организации
                re.compile(r'(?:ИНН|inn|taxpayer\s+id)[:\s=]*(\d{10,12})', re.IGNORECASE)
            ],
            'phone': [
                re.compile(r'(?:\+7|8)[\s\-]?\(?(\d{3})\)?[\s\-]?(\d{3})[\s\-]?(\d{2})[\s\-]?(\d{2})'),
                re.compile(
                    r'(?:тел|phone|моб|mobile|контакт)[.:\s]*((?:\+7|8)[\s\-]?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2})',
                    re.IGNORECASE)
            ],
            'email': [
                re.compile(r'\b([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})\b'),
                re.compile(r'(?:e-mail|email|почта)[:\s]*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
                           re.IGNORECASE)
            ],
            'fio': [
                re.compile(r'\b([А-ЯЁ][а-яё]{2,}\s+[А-ЯЁ][а-яё]{2,}(?:\s+[А-ЯЁ][а-яё]{2,})?)\b'),
                re.compile(
                    r'(?:ФИО|Ф\.И\.О\.|фио|fio|full\s+name)[:\s]*([А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+\s+[А-ЯЁ][а-яё]+)',
                    re.IGNORECASE)
            ],
            'address': [
                re.compile(r'(?:г\.|город|ул\.|улица|д\.|дом|кв\.|квартира)[\s\.]*([А-ЯЁа-яё0-9\s\-,./]+)',
                           re.IGNORECASE),
                re.compile(r'(?:адрес|address)[:\s]*([^.]+)', re.IGNORECASE)
            ],
            'birth_date': [
                re.compile(r'\b(\d{2}[./-]\d{2}[./-]\d{4})\b'),
                re.compile(r'(?:дата\s+рождения|birth\s+date|born)[:\s]*(\d{2}[./-]\d{2}[./-]\d{4})', re.IGNORECASE)
            ]
        }

        # Улучшенные паттерны для исключения
        self.exclude_patterns = [
            re.compile(r'\b(0{3,}|1{3,}|2{3,}|3{3,}|4{3,}|5{3,}|6{3,}|7{3,}|8{3,}|9{3,})\b'),
            re.compile(r'\b\d{4}[./-]\d{2}[./-]\d{2}\b'),
            re.compile(r'\b\d+\.\d+\.\d+\.\d+\b'),  # IP адреса
            re.compile(r'\b(?:http|https|ftp)://[^\s]+'),
            re.compile(r'\b[0-9A-F]{8}-[0-9A-F]{4}-[0-9A-F]{4}-[0-9A-F]{4}-[0-9A-F]{12}\b', re.I),
            re.compile(r'\b(?:ОГРН|ОКПО|КПП|БИК)[\s:]*\d+\b', re.I),
            re.compile(r'\b[A-Z]{2}\d{6,}\b'),  # Серийные номера
            re.compile(r'\b\d{5,6}\b'),  # Почтовые индексы
        ]

        self.pd_keywords = {
            'high': ['паспорт', 'passport', 'снилс', 'snils', 'инн', 'inn',
                     'персональные данные', 'personal data', 'конфиденциально'],
            'medium': ['телефон', 'phone', 'адрес', 'address', 'фио', 'fio',
                       'дата рождения', 'birth date', 'серия', 'номер'],
            'low': ['e-mail', 'email', 'почта', 'контакт', 'contact']
        }

        # Частотность слов в русском языке для фильтрации ложных ФИО
        self.common_words = {'россия', 'москва', 'российской', 'федерации', 'компания',
                             'организация', 'система', 'данные', 'документ', 'приказ'}

    def calculate_entropy(self, text):
        """Вычисление энтропии текста для определения случайных данных"""
        if not text:
            return 0
        prob = [float(text.count(c)) / len(text) for c in set(text)]
        return -sum(p * np.log2(p) for p in prob)

    def validate_snils_advanced(self, snils_str):
        """Продвинутая валидация СНИЛС"""
        try:
            digits = re.sub(r'[\s\-]', '', snils_str)
            if len(digits) != 11:
                return False, "invalid_length"

            if digits == '0' * 11:
                return False, "all_zeros"

            if len(set(digits)) < 5:  # Слишком много повторений
                return False, "low_entropy"

            control = int(digits[9:11])
            sum_val = sum(int(digits[i]) * (9 - i) for i in range(9))

            if sum_val < 100:
                calculated = sum_val
            elif sum_val in (100, 101):
                calculated = 0
            else:
                calculated = sum_val % 101
                if calculated == 100:
                    calculated = 0

            if calculated != control:
                return False, "invalid_checksum"

            return True, "valid"
        except:
            return False, "error"

    def validate_inn_advanced(self, inn_str):
        """Продвинутая валидация ИНН"""
        try:
            digits = re.sub(r'\D', '', inn_str)

            # ИНН физлица (12 цифр)
            if len(digits) == 12:
                if digits == '0' * 12:
                    return False, "all_zeros"

                coeff1 = [7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
                n1 = sum(int(digits[i]) * coeff1[i] for i in range(10)) % 11 % 10

                coeff2 = [3, 7, 2, 4, 10, 3, 5, 9, 4, 6, 8]
                n2 = sum(int(digits[i]) * coeff2[i] for i in range(11)) % 11 % 10

                if n1 == int(digits[10]) and n2 == int(digits[11]):
                    return True, "valid_12"
                else:
                    return False, "invalid_checksum"

            # ИНН организации (10 цифр)
            elif len(digits) == 10:
                coeff = [2, 4, 10, 3, 5, 9, 4, 6, 8]
                n = sum(int(digits[i]) * coeff[i] for i in range(9)) % 11 % 10

                if n == int(digits[9]):
                    return True, "valid_10"
                else:
                    return False, "invalid_checksum"

            return False, "invalid_length"
        except:
            return False, "error"

    def validate_phone_advanced(self, phone_str):
        """Продвинутая валидация телефона"""
        digits = re.sub(r'\D', '', phone_str)

        # Приводим к 11-значному формату
        if len(digits) == 10:
            digits = '8' + digits
        elif len(digits) != 11:
            return False, "invalid_length"

        # Проверка первого символа
        if digits[0] not in '789':
            return False, "invalid_first_digit"

        # Проверка на паттерны
        if re.match(r'^[0-9]*(.)\1{4,}[0-9]*$', digits):
            return False, "repeating_digits"

        if digits in ['80000000000', '84950000000', '84991234567']:
            return False, "test_number"

        # Проверка кода оператора для мобильных
        if digits[0] == '9':
            operator_code = digits[1:4]
            valid_codes = {'900', '901', '902', '903', '904', '905', '906', '908', '909',
                           '910', '911', '912', '913', '914', '915', '916', '917', '918', '919',
                           '920', '921', '922', '923', '924', '925', '926', '927', '928', '929',
                           '930', '931', '932', '933', '934', '935', '936', '937', '938', '939',
                           '950', '951', '952', '953', '954', '955', '956', '957', '958', '959',
                           '960', '961', '962', '963', '964', '965', '966', '967', '968', '969',
                           '970', '971', '972', '973', '974', '975', '976', '977', '978', '979',
                           '980', '981', '982', '983', '984', '985', '986', '987', '988', '989',
                           '991', '992', '993', '994', '995', '996', '997', '998', '999'}
            if operator_code not in valid_codes:
                return False, "invalid_operator"

        return True, "valid"

    def validate_passport_advanced(self, passport_str):
        """Продвинутая валидация паспорта"""
        digits = re.sub(r'\D', '', passport_str)

        if len(digits) != 10:
            return False, "invalid_length"

        series = digits[:4]
        number = digits[4:]

        # Проверка серии
        if series == '0000':
            return False, "invalid_series_zeros"

        # Проверка на реальные серии паспортов (упрощенно)
        valid_series_start = {'01', '02', '03', '04', '05', '06', '07', '08', '09',
                              '10', '11', '12', '13', '14', '15', '16', '17', '18', '19',
                              '20', '21', '22', '23', '24', '25', '26', '27', '28', '29',
                              '30', '31', '32', '33', '34', '35', '36', '37', '38', '39',
                              '40', '41', '42', '43', '44', '45', '46', '47', '48', '49',
                              '50', '51', '52', '53', '54', '55', '56', '57', '58', '59',
                              '60', '61', '62', '63', '64', '65', '66', '67', '68', '69',
                              '70', '71', '72', '73', '74', '75', '76', '77', '78', '79',
                              '80', '81', '82', '83', '84', '85', '86', '87', '88', '89',
                              '90', '91', '92', '93', '94', '95', '96', '97', '98', '99'}

        if series[:2] not in valid_series_start:
            return False, "invalid_series_format"

        # Проверка номера
        if number == '000000':
            return False, "invalid_number_zeros"

        if len(set(number)) < 3:  # Слишком много повторений
            return False, "low_entropy"

        return True, "valid"

    def validate_fio_advanced(self, fio_str):
        """Продвинутая валидация ФИО"""
        parts = fio_str.split()

        if len(parts) < 2 or len(parts) > 4:
            return False, "invalid_parts_count"

        for part in parts:
            # Проверка на длину
            if len(part) < 2 or len(part) > 20:
                return False, "invalid_length"

            # Проверка на заглавную букву
            if not part[0].isupper() or not part[0] in 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ':
                return False, "invalid_first_letter"

            # Проверка на строчные буквы
            if not all(c.islower() or c in '-ё' for c in part[1:]):
                return False, "invalid_case"

            # Проверка на частотные слова
            if part.lower() in self.common_words:
                return False, "common_word"

        # Проверка на реалистичность окончаний
        endings = {'ов', 'ев', 'ин', 'ын', 'ий', 'ой', 'ый', 'ая', 'яя', 'ко', 'ук', 'юк'}
        has_valid_ending = any(parts[-1].lower().endswith(ending) for ending in endings)

        return True, "valid" if has_valid_ending else "suspicious"

    def calculate_confidence_score(self, matches):
        """Расчет общей уверенности в наличии ПДн"""
        if not matches:
            return 0.0

        score = 0.0
        weights = {
            'snils': 1.0,
            'passport': 0.9,
            'inn': 0.8,
            'phone': 0.6,
            'email': 0.5,
            'fio': 0.4,
            'address': 0.3,
            'birth_date': 0.3
        }

        # Базовая оценка по типам данных
        data_types_found = set()
        for match in matches:
            data_type = match['type']
            data_types_found.add(data_type)
            score += weights.get(data_type, 0.2) * match.get('confidence', 0.5)

        # Бонус за разнообразие типов данных
        if len(data_types_found) >= 3:
            score *= 1.5
        elif len(data_types_found) >= 2:
            score *= 1.2

        # Нормализация
        return min(score / 2.0, 1.0)

    def has_personal_data(self, text):
        """Улучшенное определение ПДн с оценкой уверенности"""
        if not text or len(text) < 10:
            return False, 0.0, []

        text_lower = text.lower()

        # Оценка наличия ключевых слов
        keyword_score = 0
        for priority, keywords in self.pd_keywords.items():
            for keyword in keywords:
                if keyword in text_lower:
                    keyword_score += {'high': 0.3, 'medium': 0.2, 'low': 0.1}[priority]

        found_matches = []

        # Поиск по всем паттернам
        for data_type, patterns in self.patterns.items():
            for pattern in patterns:
                matches = pattern.finditer(text)

                for match in matches:
                    if len(match.groups()) > 1:
                        value = ''.join(g for g in match.groups() if g)
                    else:
                        value = match.group(1) if match.groups() else match.group(0)

                    if not value:
                        continue

                    # Проверка на исключения
                    is_excluded = False
                    for exclude_pattern in self.exclude_patterns:
                        if exclude_pattern.search(value):
                            is_excluded = True
                            break

                    if is_excluded:
                        continue

                    # Валидация
                    is_valid = False
                    confidence = 0.5

                    if data_type == 'snils':
                        is_valid, reason = self.validate_snils_advanced(value)
                        confidence = 0.9 if is_valid else 0.3
                    elif data_type == 'inn':
                        is_valid, reason = self.validate_inn_advanced(value)
                        confidence = 0.9 if is_valid else 0.3
                    elif data_type == 'phone':
                        is_valid, reason = self.validate_phone_advanced(value)
                        confidence = 0.8 if is_valid else 0.3
                    elif data_type == 'passport':
                        is_valid, reason = self.validate_passport_advanced(value)
                        confidence = 0.9 if is_valid else 0.3
                    elif data_type == 'fio':
                        is_valid, reason = self.validate_fio_advanced(value)
                        confidence = 0.7 if is_valid else 0.2
                    elif data_type == 'email':
                        is_valid = True
                        confidence = 0.7
                    elif data_type in ['address', 'birth_date']:
                        is_valid = True
                        confidence = 0.5

                    if is_valid:
                        # Контекстная проверка
                        context_score = self.check_context(text, value, match.start())
                        confidence *= (1 + context_score)

                        found_matches.append({
                            'type': data_type,
                            'value': value,
                            'confidence': min(confidence, 1.0),
                            'reason': reason if 'reason' in locals() else 'valid'
                        })

        # Расчет общей уверенности
        confidence_score = self.calculate_confidence_score(found_matches)

        # Учитываем ключевые слова
        confidence_score += keyword_score * 0.3

        # Принимаем решение
        threshold = 0.4
        if confidence_score >= threshold:
            return True, confidence_score, found_matches
        else:
            return False, confidence_score, found_matches

    def check_context(self, text, value, position):
        """Проверка контекста вокруг найденного значения"""
        start = max(0, position - 150)
        end = min(len(text), position + len(value) + 150)
        context = text[start:end].lower()

        context_score = 0

        # Поиск ключевых слов в контексте
        for priority, keywords in self.pd_keywords.items():
            weight = {'high': 0.3, 'medium': 0.2, 'low': 0.1}[priority]
            for keyword in keywords:
                if keyword in context:
                    context_score += weight

        # Проверка на табличную структуру (вертикальные разделители)
        if '|' in context or '\t' in context:
            context_score += 0.1

        # Проверка на наличие меток
        if ':' in context[:position] if position > 0 else False:
            context_score += 0.1

        return min(context_score, 0.5)


class EnhancedFileProcessor:
    def __init__(self, max_workers=4):
        self.detector = AdvancedPersonalDataDetector()
        self.max_workers = max_workers
        self.cache_file = "ocr_cache_enhanced.json"
        self.cache = self.load_cache()
        self.stats = defaultdict(int)

    def load_cache(self):
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return {}
        return {}

    def save_cache(self):
        try:
            if len(self.cache) > 2000:
                items = list(self.cache.items())
                self.cache = dict(sorted(items, key=lambda x: x[1].get('timestamp', 0))[-1000:])
            with open(self.cache_file, 'w', encoding='utf-8') as f:
                json.dump(self.cache, f)
        except:
            pass

    def get_file_hash(self, filepath):
        try:
            with open(filepath, 'rb') as f:
                return hashlib.md5(f.read(131072)).hexdigest()
        except:
            return None

    def find_files(self, directory):
        directory_path = Path(directory)
        extensions = {'.pdf', '.docx', '.xlsx', '.xls', '.txt', '.csv',
                      '.tif', '.tiff', '.jpg', '.jpeg', '.png', '.bmp'}

        for file_path in directory_path.rglob('*'):
            if file_path.suffix.lower() in extensions:
                if file_path.stat().st_size < 100 * 1024 * 1024:  # Пропускаем файлы > 100MB
                    yield file_path

    def preprocess_image_advanced(self, img):
        """Продвинутая предобработка изображений"""
        try:
            # Конвертация в grayscale
            if img.mode != 'L':
                img = img.convert('L')

            # Увеличение контраста
            enhancer = ImageEnhance.Contrast(img)
            img = enhancer.enhance(2.0)

            # Увеличение резкости
            enhancer = ImageEnhance.Sharpness(img)
            img = enhancer.enhance(1.5)

            # Применение фильтров
            img = img.filter(ImageFilter.MedianFilter(size=3))

            # Бинаризация
            threshold = 127
            img = img.point(lambda p: 255 if p > threshold else 0)

            # Увеличение при необходимости
            if img.width < 2000:
                ratio = 2000 / img.width
                new_size = (int(img.width * ratio), int(img.height * ratio))
                img = img.resize(new_size, Image.Resampling.LANCZOS)

            return img
        except:
            return img

    def extract_text_from_image(self, filepath, lang='rus+eng'):
        try:
            file_hash = self.get_file_hash(filepath)
            if file_hash and file_hash in self.cache:
                self.stats['cache_hits'] += 1
                return self.cache[file_hash]['text']

            with Image.open(filepath) as img:
                # Разные варианты предобработки
                processed_variants = []

                # Вариант 1: Базовая обработка
                if img.mode != 'L':
                    img_gray = img.convert('L')
                else:
                    img_gray = img
                processed_variants.append(img_gray)

                # Вариант 2: Продвинутая обработка
                processed_variants.append(self.preprocess_image_advanced(img))

                # Вариант 3: Без обработки
                processed_variants.append(img)

                texts = []
                configs = [
                    '--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ@.+-_',
                    '--oem 3 --psm 3',
                    '--oem 1 --psm 11'
                ]

                for variant in processed_variants:
                    for config in configs[:2]:
                        try:
                            text = pytesseract.image_to_string(variant, lang=lang, config=config)
                            if text.strip():
                                texts.append(text)
                        except:
                            continue

                if not texts:
                    return ""

                # Выбор лучшего результата по метрикам
                best_text = max(texts, key=lambda t: (
                        sum(c.isdigit() for c in t) * 2 +
                        sum(c.isalpha() for c in t)
                ))

                best_text = re.sub(r'\s+', ' ', best_text).strip()

                if file_hash and len(best_text) > 50:
                    self.cache[file_hash] = {
                        'text': best_text,
                        'timestamp': datetime.now().timestamp()
                    }

                return best_text

        except Exception as e:
            self.stats['image_errors'] += 1
            return ""

    def extract_text_from_pdf(self, filepath):
        try:
            import fitz
            doc = fitz.open(filepath)
            text_parts = []

            for page_num in range(min(len(doc), 50)):  # Ограничиваем до 50 страниц
                page = doc[page_num]

                # Извлечение текста
                page_text = page.get_text()

                # Извлечение текста из аннотаций и форм
                for annot in page.annots():
                    if annot.info.get('content'):
                        page_text += ' ' + annot.info['content']

                if len(page_text.strip()) < 100:
                    # OCR для страниц с малым количеством текста
                    matrix = fitz.Matrix(2.5, 2.5)  # Увеличенное разрешение
                    pix = page.get_pixmap(matrix=matrix)
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                    # Обработка частями для больших изображений
                    if img.width > 3000 or img.height > 3000:
                        # Разбиваем на части
                        tile_size = 2000
                        for y in range(0, img.height, tile_size):
                            for x in range(0, img.width, tile_size):
                                tile = img.crop((x, y, min(x + tile_size, img.width),
                                                 min(y + tile_size, img.height)))
                                ocr_text = self.extract_text_from_image_obj(tile)
                                if ocr_text:
                                    text_parts.append(ocr_text)
                    else:
                        ocr_text = self.extract_text_from_image_obj(img)
                        if ocr_text:
                            text_parts.append(ocr_text)
                else:
                    text_parts.append(page_text)

                if len(' '.join(text_parts)) > 50000:
                    break

            doc.close()
            return ' '.join(text_parts)

        except Exception as e:
            self.stats['pdf_errors'] += 1
            return ""

    def extract_text_from_image_obj(self, img):
        """OCR для объекта PIL Image"""
        try:
            processed = self.preprocess_image_advanced(img)
            text = pytesseract.image_to_string(processed, lang='rus+eng', config='--oem 3 --psm 6')
            return re.sub(r'\s+', ' ', text).strip()
        except:
            return ""

    def extract_text_from_docx(self, filepath):
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []

            # Текст из параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))

            # Текст из колонтитулов
            for section in doc.sections:
                for header in section.header.paragraphs:
                    if header.text.strip():
                        text_parts.append(header.text)
                for footer in section.footer.paragraphs:
                    if footer.text.strip():
                        text_parts.append(footer.text)

            # Текст из текстовых полей
            try:
                for shape in doc.inline_shapes:
                    if shape.type == 3:  # Textbox
                        if hasattr(shape, 'text'):
                            text_parts.append(shape.text)
            except:
                pass

            return ' '.join(text_parts)

        except Exception as e:
            self.stats['docx_errors'] += 1
            return ""

    def extract_text_from_excel(self, filepath):
        try:
            import pandas as pd
            text_parts = []

            excel_file = pd.ExcelFile(filepath)
            for sheet_name in excel_file.sheet_names:
                # Чтение с обработкой разных типов данных
                df = pd.read_excel(filepath, sheet_name=sheet_name, nrows=5000, dtype=str)

                # Поиск в заголовках
                text_parts.append(f"Sheet: {sheet_name}")
                text_parts.append(f"Columns: {' | '.join(df.columns)}")

                # Анализ данных
                for col in df.columns:
                    # Пропускаем пустые колонки
                    if df[col].isna().all():
                        continue

                    col_data = df[col].dropna()
                    if col_data.empty:
                        continue

                    # Ограничиваем количество значений
                    values = col_data.head(200).tolist()
                    if values:
                        text_parts.append(f"{col}: {' '.join(str(v) for v in values)}")

                if len(' '.join(text_parts)) > 100000:
                    break

            return ' '.join(text_parts)

        except Exception as e:
            self.stats['excel_errors'] += 1
            return ""

    def extract_text_from_txt(self, filepath):
        """Извлечение текста из текстовых файлов"""
        encodings = ['utf-8', 'windows-1251', 'cp866', 'koi8-r']

        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding) as f:
                    return f.read(100000)  # Ограничиваем объем
            except:
                continue
        return ""

    def process_file(self, filepath):
        try:
            file_path_obj = Path(filepath)
            file_ext = file_path_obj.suffix.lower()
            file_size = os.path.getsize(filepath)

            # Извлечение текста
            if file_ext == '.pdf':
                text = self.extract_text_from_pdf(filepath)
            elif file_ext == '.docx':
                text = self.extract_text_from_docx(filepath)
            elif file_ext in ['.xlsx', '.xls']:
                text = self.extract_text_from_excel(filepath)
            elif file_ext in ['.txt', '.csv']:
                text = self.extract_text_from_txt(filepath)
            elif file_ext in ['.tif', '.tiff', '.jpg', '.jpeg', '.png', '.bmp']:
                text = self.extract_text_from_image(filepath)
            else:
                return None

            if text and len(text) > 20:
                has_pd, confidence, matches = self.detector.has_personal_data(text)

                if has_pd:
                    mod_time = datetime.fromtimestamp(os.path.getmtime(filepath))

                    # Определяем типы найденных данных
                    data_types = list(set(m['type'] for m in matches))

                    return {
                        'size': self.format_size(file_size),
                        'time': mod_time.strftime("%Y-%m-%d %H:%M:%S"),
                        'name': file_path_obj.name,
                        'full_path': str(filepath),
                        'extension': file_ext,
                        'confidence': f"{confidence:.2%}",
                        'data_types': ', '.join(data_types),
                        'matches_count': len(matches)
                    }

            self.stats['processed'] += 1
            return None

        except Exception as e:
            self.stats['file_errors'] += 1
            return None

    def format_size(self, size):
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"

    def process_directory(input_dir, output_csv):
        files = list(self.find_files(input_dir))
        if not files:
            print("No supported files found")
            return

        total = len(files)
        print(f"Found {total} files, processing with {self.max_workers} workers...")

        results = []
        processed = 0
        found_count = 0

        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            future_to_file = {executor.submit(self.process_file, str(f)): f for f in files}

            for future in as_completed(future_to_file):
                processed += 1
                result = future.result()

                if result:
                    results.append(result)
                    found_count += 1
                    print(f"✓ [{processed}/{total}] Found PD in {result['name']} (confidence: {result['confidence']})")
                else:
                    if processed % 20 == 0:
                        print(f"Progress: {processed}/{total} files ({found_count} with PD found)")

        self.save_cache()

        if results:
            # Сортировка по уверенности
            results.sort(key=lambda x: float(x['confidence'].rstrip('%')) / 100, reverse=True)

            fieldnames = ['size', 'time', 'name', 'full_path', 'extension',
                          'confidence', 'data_types', 'matches_count']

            with open(output_csv, 'w', newline='', encoding='utf-8-sig') as csvfile:
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()
                writer.writerows(results)

        print(f"\n{'=' * 60}")
        print(f"COMPLETED:")
        print(f"  Total files: {total}")
        print(f"  PD found: {found_count}")
        print(f"  Cache hits: {self.stats.get('cache_hits', 0)}")
        print(f"  Results saved to: {output_csv}")
        print(f"{'=' * 60}")


if __name__ == "__main__":
    input_dir = "./"
    output_csv = "result.csv"

    processor = EnhancedFileProcessor(max_workers=4)
    processor.process_directory(input_dir, output_csv)